"""ZONE 2a — text acquisition + OCR.

`get_document_text` turns a DiscoveredDoc into raw text plus OCRMetrics,
branching by format:
  • html       → strip tags (BeautifulSoup) to readable text
  • pdf_text   → pdfplumber/pypdf text layer
  • pdf_scanned→ pluggable OCR provider (mock|tesseract|paddle|azure)

For text PDFs we also detect a "secretly scanned" PDF (empty text layer) and fall
back to OCR automatically — a common real-world failure mode.
"""
from __future__ import annotations

import time as _time

from .. import metering as _metering

import hashlib
import json
import re
from pathlib import Path

from ..config import settings
from ..providers import get_ocr_provider
from ..providers.ocr_base import OCRProvider
from ..schemas import DiscoveredDoc, DocFormat, OCRMetrics


#: Class and id fragments that mark site chrome. The semantic tags below cover a page written
#: after about 2015; government portals frequently are not. cac.gov.cn closes every page with
#: a <div class="footer"> carrying the ministry name, an ICP registration number, a WeChat
#: link and 返回顶部 — and because that div is not a <footer>, all of it survived and was
#: absorbed by the LAST provision, whose body runs to the end of the text. PIPL article 74
#: ("this Law takes effect on 1 November 2021") came out with 200 characters of copyright
#: notice attached to it, in the Verbatim Snippet column.
_CHROME = ("footer", "copyright", "beian", "navbar", "nav-", "menu", "breadcrumb", "sidebar",
           "side-bar", "share", "related", "totop", "back-to-top", "banner", "search-box")


def _html_to_text(html: str) -> str:
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "nav", "header", "footer", "aside", "noscript"]):
        tag.decompose()

    # Then the same thing by class and id. The size guard matters: a portal that calls its
    # main content div "page-menu" would otherwise lose the entire document, and a rule that
    # can delete the law is worse than a footer that survives. Anything holding more than a
    # fifth of the page is treated as content whatever it calls itself.
    whole = max(len(soup.get_text(strip=True)), 1)
    for tag in soup.find_all(attrs={"class": True}) + soup.find_all(attrs={"id": True}):
        if tag.decomposed:
            continue
        ident = " ".join(tag.get("class") or []) + " " + (tag.get("id") or "")
        if not any(word in ident.lower() for word in _CHROME):
            continue
        if len(tag.get_text(strip=True)) < whole * 0.2:
            tag.decompose()
    return soup.get_text("\n", strip=True)


# Single-page-app framework markers. A page carrying one of these renders its body with
# JavaScript, so a static fetch returns only the app shell (site chrome, no law text).
# legislation.gov.au (Angular: `ng-version`) is the Round-1 case.
_SPA_MARKERS = ("ng-version=", "<app-root", "data-reactroot", "__next_data__",
                "window.__nuxt", 'id="__nuxt"',
                # Vite build, which is what China's flk.npc.gov.cn ships: the whole page is
                # <div id="app"></div> plus one module bundle, and de-chroming it yielded the
                # 9-character site title, which then became a "provision" citing nothing.
                'src="/assets/index-', '<div id="app"></div>')


def is_js_app_shell(html: str, text: str | None = None) -> bool:
    """True when `html` is an UNRENDERED SPA shell — a JS framework marker is present AND
    the de-chromed text has no legal structure (so extraction would otherwise emit the
    site's navigation chrome as a bogus, non-verbatim 'provision'). If real section
    markers ARE present the page carries server-rendered content and is kept."""
    head = html[:400_000].lower()
    if not any(m in head for m in _SPA_MARKERS):
        return False
    from .extraction import SECTION_RE, _STRUCT_RE_CN, _STRUCT_RE_MN
    body = text if text is not None else _html_to_text(html)
    # "Has legal structure" has to be asked in every drafting convention we support. SECTION_RE
    # is Latin-only, so a genuinely server-rendered Chinese or Mongolian statute matches none of
    # it — and would be thrown away as an empty shell, losing the whole document.
    return not any(rx.search(body) for rx in (SECTION_RE, _STRUCT_RE_CN, _STRUCT_RE_MN))


def _page_count(path: str) -> int:
    try:
        import pypdfium2
        doc = pypdfium2.PdfDocument(path)
        try:
            return len(doc)
        finally:
            doc.close()
    except Exception:
        try:
            from pypdf import PdfReader
            return len(PdfReader(path).pages)
        except Exception:
            return 1


# A bold, number-led line is a section heading. Some PDFs (AU legislation.gov.au) number
# sections "77 Requirement…" with NO keyword and NO period — invisible to SECTION_RE — but
# they ARE typeset bold/larger than the body, so the font is the reliable signal. We mark
# such lines with a record-separator (\x1e, never present in legal text) so extraction can
# split on them; SG/MY PDFs don't bold their headings, so nothing is marked and they fall
# back to the regex path unchanged.
HEADING_MARK = "\x1e"
_HEADING_NUM_RE = re.compile(r"^\s*\d{1,3}[A-Za-z]{0,2}\s+[A-Za-z]")

# Page-boundary sentinel. The submission template's "Location Reference" column is a page
# number, and it used to be INTERPOLATED from the character offset (offset/total * pages) —
# which assumes every page holds the same number of characters. Real statute PDFs are nothing
# like that: schedules, tables and forms swing density wildly, so on a 1,800-page consolidated
# Act the cited page could be tens of pages out. An audit sampling provisions and re-reading
# the cited page with a second extractor found the citation wrong about half the time, while
# the snippet text itself was perfectly verbatim — a defect no amount of reading the output
# would reveal, because the text looks right.
# Fix: mark the real page boundaries in the text itself. \x0c (form feed) is the conventional
# page separator, never appears in legal drafting, and travels WITH the text through every
# later chrome/TOC strip — so counting sentinels before an offset gives the true page even
# after the text has been rewritten. A newline follows it so line-anchored section regexes
# still match at a page start.
#
# The sentinel carries its ABSOLUTE page number (`\x0c12\x0c`) rather than being a bare
# delimiter. Counting bare delimiters looked sufficient and was not: `_strip_arrangement_toc`
# and the chrome strippers DELETE whole spans of text, taking any delimiters inside them with
# it, so every page after a deleted block was numbered short. A self-describing sentinel
# survives arbitrary deletion — the last one before an offset still states its own page.
# It occupies its own line so the line-anchored section regexes cannot match it.
PAGE_MARK = "\x0c"
PAGE_MARK_RE = re.compile(r"\x0c(\d+)\x0c")


def _join_pages(pages: list[str]) -> str:
    """Join extracted pages, tagging each with a self-describing page sentinel."""
    return "\n\n".join(f"{PAGE_MARK}{i}{PAGE_MARK}\n{p}"
                       for i, p in enumerate(pages, start=1))


def _is_bold_heading(text: str, chars: list) -> bool:
    if not _HEADING_NUM_RE.match(text):
        return False
    glyphs = [c for c in (chars or []) if (c.get("text") or "").strip()]
    if not glyphs:
        return False
    bold = sum(1 for c in glyphs if "bold" in (c.get("fontname") or "").lower())
    return bold / len(glyphs) >= 0.6


def _strip_running_chrome(pages: list[str]) -> list[str]:
    """General page running-header/footer removal — the robust alternative to per-law patterns.

    A running header/footer (the act title + page number, "Section 77A" / "Division 1" banners,
    "S 63/2021 14", "33 Act 2012 2020 Ed.") sits in the TOP or BOTTOM band of the page and
    repeats across pages; only its page number changes. So: mask digit runs (page numbers),
    look at the first 2 / last 3 lines of every page, and treat any normalised line that recurs
    in that band on a large fraction of pages as chrome — then drop those lines everywhere.
    Body text and one-off marginal headings don't repeat in the band, so they're kept."""
    if len(pages) < 3:
        return pages
    from collections import Counter

    def _norm(s: str) -> str:
        return re.sub(r"\d+", "#", s.replace(HEADING_MARK, "")).strip()

    counts: Counter = Counter()
    for pg in pages:
        ls = [l for l in pg.split("\n") if l.strip()]
        for l in ls[:2] + ls[-3:]:                     # header band + footer band
            n = _norm(l)
            if 0 < len(n) <= 80:
                counts[n] += 1
    thresh = max(3, int(0.25 * len(pages)))            # recurs on ≥¼ of pages → chrome
    chrome = {n for n, c in counts.items() if c >= thresh}
    if not chrome:
        return pages
    return ["\n".join(l for l in pg.split("\n") if _norm(l) not in chrome) for pg in pages]


def _pdf_page_texts(path: str, pages: list[int] | None = None) -> dict[int, str]:
    """{1-indexed page: text-layer text} for the requested pages (all pages when None).

    Returned per page rather than pre-joined so the caller can (a) run the density check page
    by page and (b) splice OCR output into the gaps of a mixed document. Chrome stripping is
    deliberately NOT applied here: it needs the whole page set to spot what repeats.
    """
    want = set(pages) if pages else None
    out: dict[int, str] = {}
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for idx, pg in enumerate(pdf.pages, start=1):
                if want is not None and idx not in want:
                    continue
                try:
                    # x_tolerance=2 infers spaces from glyph gaps — without it legal PDFs come
                    # out jammed ("personaldatamustnot"); matches the old extract_text() spacing.
                    lines = pg.extract_text_lines(x_tolerance=2)   # per-line text + char fonts
                except Exception:
                    lines = None
                if lines:
                    out[idx] = "\n".join(
                        (HEADING_MARK + ln["text"]) if _is_bold_heading(ln["text"], ln.get("chars"))
                        else ln["text"] for ln in lines)
                else:
                    out[idx] = pg.extract_text(x_tolerance=2, y_tolerance=3) or ""
        return out
    except Exception:
        pass
    try:
        from pypdf import PdfReader
        for idx, pg in enumerate(PdfReader(path).pages, start=1):
            if want is not None and idx not in want:
                continue
            out[idx] = pg.extract_text() or ""
    except Exception:
        return out
    return out


def _pdf_text_layer(path: str) -> str:
    # x_tolerance infers spaces from glyph gaps — without it, legal PDFs come out with
    # words jammed together ("Anorganisationmustnot"), which wrecks the Character Error
    # Rate and the downstream matching. pdfplumber preserves spacing + line structure.
    texts = _pdf_page_texts(path)
    if not texts:
        return ""
    ordered = [texts[k] for k in sorted(texts)]
    return _join_pages(_strip_running_chrome(ordered))


def _content_key(doc: DiscoveredDoc) -> str | None:
    """Stable hash of a document's actual BYTES (not its doc_id/title, which can differ across
    two discoveries of the identical file). Identical bytes always extract to identical text,
    so this is the cache key for get_document_text — independent of filename convention."""
    path = doc.local_path
    try:
        if path and Path(path).exists():
            data = Path(path).read_bytes()
        elif doc.raw_text:
            data = doc.raw_text.encode("utf-8", errors="ignore")
        else:
            return None
    except Exception:
        return None
    return hashlib.sha256(data).hexdigest()[:20]


# Bump when the canonical text format changes (page routing, sentinel rules, markdown
# normalisation). The cache key is content-hash + engine, so without this a format change
# would keep serving text produced by the OLD rules until the directory was purged by hand.
# v3: pages are now joined with a PAGE_MARK sentinel so Location Reference page numbers are
# COUNTED rather than interpolated. Text cached under v2 has no marks and would silently keep
# producing estimated page citations.
EXTRACT_FORMAT_VERSION = "v4"


def _extract_cache_path(key: str, provider_name: str) -> Path:
    d = settings.cache_path / "_extracted"
    d.mkdir(parents=True, exist_ok=True)
    return d / f"{key}_{provider_name}_{EXTRACT_FORMAT_VERSION}.json"


def get_document_text(doc: DiscoveredDoc, ocr_provider: OCRProvider | None = None) -> tuple[str, OCRMetrics]:
    """Cached wrapper: extraction (OCR/PDF-to-text) is the single biggest per-run cost bucket
    (~44% of wall-clock on a live crawl — bigger than embedding + LLM grading combined), yet the
    SAME document (same bytes, same OCR provider) always extracts to the SAME text. A repeat run
    within fetch_ttl_hours, or the same law showing up for both Pillar 6 and Pillar 7, would
    otherwise re-run pdfplumber/MarkItDown/OCR on it every time. Cache the (text, metrics) result
    keyed by content hash + provider; the fetched BODY cache (fetch.py) already makes this safe —
    this just avoids re-parsing bytes we've already parsed."""
    if not settings.extraction_cache_enabled:
        return _extract_document_text(doc, ocr_provider)
    provider_name = (ocr_provider or get_ocr_provider(settings.ocr_provider)).name
    key = _content_key(doc)
    if key is None:
        return _extract_document_text(doc, ocr_provider)
    cache_f = _extract_cache_path(key, provider_name)
    if cache_f.exists():
        try:
            data = json.loads(cache_f.read_text(encoding="utf-8"))
            return data["text"], OCRMetrics.model_validate(data["metrics"])
        except Exception:
            pass  # unreadable/stale cache -> fall through and re-extract
    text, metrics = _extract_document_text(doc, ocr_provider)
    try:
        cache_f.write_text(json.dumps({"text": text, "metrics": metrics.model_dump()}), encoding="utf-8")
    except Exception:
        pass  # caching is best-effort; never fail the run over a write error
    return text, metrics


def _word_to_text(path: str, metrics: OCRMetrics) -> str:
    """Text of a WORD document, via MarkItDown then python-docx.

    Returning "" on failure is the point: the previous behaviour read the file as plain text,
    and since .docx is a ZIP container the pipeline happily produced a provision whose
    verbatim snippet began "PK docProps/app.xml". Empty text yields no provisions, which is
    a visible absence rather than a citation of binary noise.
    """
    try:
        from markitdown import MarkItDown
        text = (MarkItDown().convert(path).text_content or "").strip()
        if text:
            metrics.provider = "markitdown"
            return markdown_to_plain(text)
    except Exception:
        pass
    try:
        import docx                                  # python-docx
        paragraphs = [p.text for p in docx.Document(path).paragraphs]
        text = "\n".join(paragraphs).strip()
        if text:
            metrics.provider = "python-docx"
            return text
    except Exception:
        pass
    metrics.notes = "word_unreadable"
    return ""


def _extract_document_text(doc: DiscoveredDoc, ocr_provider: OCRProvider | None = None) -> tuple[str, OCRMetrics]:
    metrics = OCRMetrics()
    path = doc.local_path
    fmt = doc.fmt

    if fmt == DocFormat.HTML or (path and path.endswith(".html")):
        raw = Path(path).read_text(encoding="utf-8", errors="ignore") if path and Path(path).exists() else (doc.raw_text or "")
        text = _html_to_text(raw)
        # An unrendered SPA shell (e.g. legislation.gov.au) carries only site chrome — return
        # empty so extraction emits no provisions instead of mapping navigation text as a law.
        if is_js_app_shell(raw, text):
            metrics.notes = "js_app_shell"
            return "", metrics
        return text, metrics

    if path and path.lower().endswith((".docx", ".doc")):
        text = _word_to_text(path, metrics)
        return to_canonical(text), metrics

    if fmt == DocFormat.TEXT or (path and path.endswith(".txt")):
        text = Path(path).read_text(encoding="utf-8", errors="ignore") if path and Path(path).exists() else (doc.raw_text or "")
        return text, metrics

    if fmt in (DocFormat.PDF_TEXT, DocFormat.PDF_SCANNED) and path:
        provider = ocr_provider or get_ocr_provider(settings.ocr_provider)
        return _extract_pdf(path, provider, metrics)

    return doc.raw_text or "", metrics


def _extract_pdf(path: str, provider, metrics: OCRMetrics) -> tuple[str, OCRMetrics]:
    """Route each page to the cheapest reader that can actually read it.

    Three outcomes, decided by `pdf_inspect.profile_pdf`:
      * no page needs OCR  → text layer only (pdfplumber: best spacing, lowest CER, no models)
      * every page needs it → the configured OCR engine over the whole file
      * some pages need it  → HYBRID, the case the old whole-file density rule could not express
    """
    from . import pdf_inspect

    prof = pdf_inspect.profile_pdf(path)
    metrics.pages = prof.page_count or _page_count(path)
    metrics.notes = (f"{metrics.notes + '; ' if metrics.notes else ''}"
                     f"triage={prof.engine}:{prof.doc_type}")

    if not prof.needs_any_ocr:
        return _pdf_text_layer(path), metrics

    if prof.is_fully_scanned or not prof.text_pages():
        text, m = _run_provider(provider, path)
        m.notes = (f"{m.notes + '; ' if m.notes else ''}triage={prof.engine}:{prof.doc_type}")
        return text, m

    # ── hybrid: text layer for the readable pages, OCR only for the scanned ones ──────────
    ocr_pages = sorted(prof.pages_needing_ocr)
    page_text = _pdf_page_texts(path, prof.text_pages())
    ocr_text, ocr_metrics = _run_provider(provider, path, pages=ocr_pages)
    for page, chunk in zip(ocr_pages, _split_pages(ocr_text, len(ocr_pages))):
        page_text[page] = chunk

    ordered = [page_text.get(p, "") for p in range(1, (prof.page_count or 0) + 1)]
    merged = _join_pages(_strip_running_chrome(ordered))
    metrics.used = True
    metrics.provider = ocr_metrics.provider
    metrics.mean_confidence = ocr_metrics.mean_confidence
    metrics.low_conf_pages = ocr_metrics.low_conf_pages
    metrics.chars = len(merged)
    metrics.notes += f"; hybrid ocr_pages={len(ocr_pages)}/{prof.page_count}"
    return merged, metrics


def _split_pages(text: str, expected: int) -> list[str]:
    """Split a provider's page-joined output back into per-page chunks.

    Providers join pages with a blank line (`OCRResult.text`). When the split does not yield
    the expected count we return the whole block as the first chunk rather than mis-aligning
    pages, so a provider that ignores the page filter degrades to "all OCR text lands in the
    first scanned slot" instead of scrambling page order.
    """
    parts = text.split("\n\n") if text else []
    return parts if len(parts) == expected else ([text] + [""] * (expected - 1) if expected else [])


# ── Canonical extraction format ───────────────────────────────────────────────────────────
# ONE format flows out of Zone 2a: line-structured PLAIN TEXT, pages joined by a blank line
# plus a PAGE_MARK sentinel (so the true page of any offset is recoverable — see PAGE_MARK),
# bold section headings prefixed with HEADING_MARK. Everything downstream is built on it —
# the per-country boundary regexes are line-anchored, `Provision.char_span` indexes into this
# exact string, and `confidence.snippet_grounding` requires the verbatim snippet to be a
# byte-substring of it. Markdown cannot be the canonical form: its syntax breaks the line
# anchors, and escaping inside a quoted provision would violate the verbatim requirement the
# rubric scores.
#
# Engines stay swappable (a judge may select MarkItDown), so any engine that emits markdown is
# normalised back to plain text HERE, at the boundary, rather than leaving the splitter to cope
# with a format it has no awareness of. Before this existed, selecting the default engine on a
# scanned PDF fed raw markdown into a plain-text-only splitter.
_MD_FENCE_RE = re.compile(r"(?m)^\s*```.*$")
_MD_HEADING_RE = re.compile(r"(?m)^\s{0,3}#{1,6}\s+")
_MD_QUOTE_RE = re.compile(r"(?m)^\s{0,3}>\s?")
_MD_HR_RE = re.compile(r"(?m)^\s{0,3}(?:[-*_]\s*){3,}$")
_MD_EMPH_RE = re.compile(r"(\*{1,3}|_{1,3})(?=\S)(.+?)(?<=\S)\1", re.S)
_MD_LINK_RE = re.compile(r"\[([^\]]*)\]\((?:[^)]*)\)")
_MD_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\((?:[^)]*)\)")
# NOTE: these use [ \t] rather than \s for the line-edge padding. \s matches \n, so a greedy
# `\s*$` swallows the newline and welds consecutive table rows into one line.
_MD_TABLE_SEP_RE = re.compile(r"(?m)^[ \t]*\|?[ \t:|-]*-{2,}[ \t:|-]*\|?[ \t]*$")


def markdown_to_plain(text: str) -> str:
    """Reduce markdown to the canonical plain-text form, preserving wording exactly.

    Only syntax is removed; no word is rewritten, reordered or re-wrapped, so a provision
    quoted from the result still matches the source document. Table rows lose their pipes and
    become space-separated cells, which is what the pdfplumber path already produces.
    """
    if not text:
        return text
    t = _MD_IMAGE_RE.sub(r"\1", text)
    t = _MD_LINK_RE.sub(r"\1", t)
    t = _MD_FENCE_RE.sub("", t)
    t = _MD_TABLE_SEP_RE.sub("", t)
    t = _MD_HR_RE.sub("", t)
    t = _MD_HEADING_RE.sub("", t)
    t = _MD_QUOTE_RE.sub("", t)
    t = _MD_EMPH_RE.sub(r"\2", t)
    # table rows: "| a | b |" → "a  b"
    t = re.sub(r"(?m)^[ \t]*\|(.+?)\|[ \t]*$",
               lambda m: "  ".join(c.strip() for c in m.group(1).split("|")), t)
    t = t.replace("\\*", "*").replace("\\_", "_").replace("\\#", "#")
    return re.sub(r"\n{3,}", "\n\n", t)


def _looks_like_markdown(text: str) -> bool:
    """Cheap positive test — only pays the normalisation cost when syntax is actually present."""
    head = text[:20000]
    return bool(_MD_HEADING_RE.search(head) or _MD_TABLE_SEP_RE.search(head)
                or _MD_IMAGE_RE.search(head) or _MD_EMPH_RE.search(head))


def to_canonical(text: str) -> str:
    """Force any engine's output into the canonical plain-text contract."""
    return markdown_to_plain(text) if text and _looks_like_markdown(text) else text


# Engines that re-read a ground-truth sidecar instead of doing raster OCR — measuring
# CER against that same sidecar would be circular (always ~0), so we skip it for them.
_SIDECAR_READERS = {"markitdown", "mock"}


def _measure_cer(provider_name: str, path: str, ocr_text: str) -> float | None:
    """Genuine Character Error Rate against a ground-truth `*.ocr.txt` sidecar, when one
    ships next to the sample (offline accuracy proof for the rubric's CER < 5% bar).
    Only meaningful for true raster-OCR engines — sidecar readers would score a fake 0."""
    if provider_name in _SIDECAR_READERS:
        return None
    ref = Path(path).with_suffix(".ocr.txt")
    if not ref.exists():
        return None
    from .cer import character_error_rate
    return character_error_rate(ref.read_text(encoding="utf-8", errors="ignore"), ocr_text)


def _run_provider(provider, path: str, pages: list[int] | None = None) -> tuple[str, OCRMetrics]:
    """Run an OCR engine and force its output into the canonical plain-text contract.

    `pages` (1-indexed) is passed through to engines that support page selection; ones that
    don't simply OCR the whole file, which costs time but never changes the result.
    """
    try:
        _t0 = _time.monotonic()
        result = provider.ocr_pdf(path, pages=pages) if pages else provider.ocr_pdf(path)
        _metering.record_ocr(provider.name, len(result.pages), _time.monotonic() - _t0)
    except TypeError:
        _t0 = _time.monotonic()
        result = provider.ocr_pdf(path)          # provider predates the `pages` argument
        _metering.record_ocr(provider.name, len(result.pages), _time.monotonic() - _t0)
    text = to_canonical(result.text)
    metrics = OCRMetrics(
        used=True,
        provider=result.provider,
        mean_confidence=result.mean_confidence,
        pages=len(result.pages),
        chars=len(text),
        low_conf_pages=result.low_conf_pages,
        cer=_measure_cer(result.provider, path, text),
    )
    return text, metrics
