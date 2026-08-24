"""The 15 October live stress test — built against the organisers' own template, not a guess.

An earlier version of this file was written from inference and got three things wrong. The
sources that settle it are in the repo: `Finalist Orientation/Live_Test_Short_Note_TEMPLATE.docx`
and `Finalist Orientation/Meeting notes.docx`. What they actually require:

  * **Any listed country, any pillar**, revealed at 10:00, about sixty minutes to finish.
    "Prepare all economies and languages in advance."
  * **Two engines**, exporting comparable results in the same format.
  * **"Engine swap should be UI-driven, not code-level."** So the switch lives on this screen.
    Reading the declaration back and calling it read-only — which is what this file used to do
    — fails the criterion it was written to satisfy.
  * **The second pass must not fetch.** The template's comparison table has a cell reading
    *"Documents fetched during this pass — must be 0"* for engine B, and the notes say it twice
    more: "the tool must be able to reprocess already-downloaded documents without re-fetching".
    Two live crawls minutes apart differ in what the portal served, so without this an engine
    comparison measures the weather.
  * **"Show cost per run and cost difference between engines."** A number each is not enough;
    the difference is the thing being asked for.
  * **Observers watch, and most actions are expected to be performed through the interface.**
    So the evidence files are downloaded here, from the run in front of you — not found later
    in an output directory.

Design notes worth keeping:

* **A step indicator, not tabs.** "Show progress for multi-step processes" (ui-ux-pro-max, ux /
  Feedback / Progress Indicators). Under a clock the operator needs to know what is done and
  what is next at a glance; tabs invite wandering.
* **Timing, cost and counts are captured, never typed.** All of them already exist on
  `RunMeta` — asking a human to copy them under time pressure only adds errors the code cannot
  make. Section 6 of the note ("anything done by hand") can then honestly say *nothing was*.
* **The comparison marks the winner per row.** Two columns of numbers make the reader do the
  arithmetic; the template asks for a difference, so the screen states one.
"""
from __future__ import annotations

import io
from datetime import datetime, timezone

import streamlit as st

from backend.config import settings
from backend.providers import registry as reg

from .home import MEASURED_PILLARS, PILLAR_SHORT

STEPS = [
    ("Brief", "Any economy, any pillar"),
    ("Run", "Engine A live · engine B on the same documents"),
    ("Compare", "Same work, two engines"),
    ("Hand in", "Evidence, comparison, note"),
]

#: The nine the panel's 2025 database covers, listed FIRST in the picker because the brief says
#: the assignment draws from them. An ordering, not a restriction — the instruction on the day
#: is "any listed country", and a picker that could not accept one would fail at the only
#: moment it exists for.
LIVE_TEST_ORDER = ("TH", "VN", "ID", "CN", "IN", "KZ", "LA", "MN", "RU")

#: Which way is better, per comparison row. Used to mark the winner rather than leaving the
#: reader to work out whether more minutes is good news.
BETTER = {"Provisions exported": "high", "Rows exported": "high",
          "Absent from the 2025 baseline": "high", "Rows needing review": "low",
          "Elapsed (minutes)": "low", "Cost of this pass (US$)": "low",
          "Found only by this engine": "high",
          "Documents fetched during this pass": None}


def new_state() -> dict:
    return {
        "step": 0,
        "brief": {},
        "runs": {},
        "engines": {
            "A": {"provider": settings.declared_engine_a_provider,
                  "model": settings.declared_engine_a_model},
            "B": {"provider": settings.declared_engine_b_provider,
                  "model": settings.declared_engine_b_model},
        },
        "notes": {"worked": "", "broke": "", "caution": "", "by_hand": ""},
        "started": None,
    }


def economy_order(codes) -> list[str]:
    """Every declared economy, the live-test nine first, then the rest."""
    nine = [c for c in LIVE_TEST_ORDER if c in codes]
    return nine + [c for c in codes if c not in nine]


# ── step indicator ───────────────────────────────────────────────────────────────────
def _steps_html(current: int) -> str:
    cells = []
    for i, (name, hint) in enumerate(STEPS):
        state = "done" if i < current else ("now" if i == current else "todo")
        mark = "✓" if i < current else str(i + 1)
        cells.append(
            f'<li class="lt-step lt-{state}">'
            f'<span class="lt-dot" aria-hidden="true">{mark}</span>'
            f'<span class="lt-name">{name}</span>'
            f'<span class="lt-hint">{hint}</span></li>')
    return (f'<ol class="lt-steps" aria-label="Live test progress">{"".join(cells)}</ol>'
            f'<p class="lt-sr">Step {current + 1} of {len(STEPS)}: {STEPS[current][0]}</p>')


CSS = """
/* !important on the layout properties, and only those. Streamlit styles `ol`/`li` inside its
   markdown container with a more specific selector than a bare class, so `display:flex` lost
   and the four steps stacked into a narrow vertical column — the border and the numbered dots
   applied, which made it look designed rather than broken. */
.lt-steps{display:flex !important;width:100%;gap:0;list-style:none !important;
  padding:0 !important;margin:0 0 1.4rem 0;
  border:1px solid var(--line);border-radius:10px;overflow:hidden}
.lt-step{flex:1 1 0;display:flex !important;flex-direction:column;gap:.15rem;
  padding:.65rem .85rem;border-right:1px solid var(--line);min-width:0;
  margin:0 !important;list-style:none !important}
.lt-step::marker{content:none}
.lt-step:last-child{border-right:0}
.lt-dot{display:inline-flex;align-items:center;justify-content:center;
  width:1.35rem;height:1.35rem;border-radius:50%;font-size:.76rem;font-weight:600;
  font-variant-numeric:tabular-nums}
.lt-name{font-weight:600;font-size:.92rem}
.lt-hint{font-size:.76rem;color:var(--muted);overflow-wrap:anywhere}
.lt-todo .lt-dot{background:var(--surface-2);color:var(--muted)}
.lt-todo .lt-name{color:var(--muted)}
.lt-now{background:color-mix(in srgb,var(--accent) 8%,transparent)}
.lt-now .lt-dot{background:var(--accent);color:#fff}
.lt-done .lt-dot{background:var(--surface-2);color:var(--accent)}
.lt-sr{position:absolute;width:1px;height:1px;overflow:hidden;clip:rect(0 0 0 0);
  white-space:nowrap}

.lt-slot{border:1px solid var(--line);border-radius:10px;padding:.8rem .95rem;
  margin-bottom:.55rem}
.lt-slot h4{margin:0 0 .3rem 0;font-size:.92rem;display:flex;align-items:baseline;gap:.45rem;
  flex-wrap:wrap}
.lt-slot h4 .eng{font-family:var(--mono);font-size:.72rem;font-weight:500;color:var(--muted);
  overflow-wrap:anywhere}
.lt-kv{display:grid;grid-template-columns:auto 1fr;gap:.12rem .9rem;font-size:.86rem;margin:0}
.lt-kv dt{color:var(--muted)}
.lt-kv dd{margin:0;font-family:var(--mono);font-variant-numeric:tabular-nums}
.lt-empty{color:var(--muted);font-size:.86rem}

/* the comparison: the winner per row is marked, because the template asks for a DIFFERENCE */
.lt-cmp{width:100%;border-collapse:collapse;font-size:.88rem}
.lt-cmp th,.lt-cmp td{padding:.42rem .6rem;border-bottom:1px solid var(--line);text-align:right}
.lt-cmp th:first-child,.lt-cmp td:first-child{text-align:left;color:var(--muted)}
.lt-cmp thead th{color:var(--ink);font-weight:600;border-bottom:2px solid var(--line)}
.lt-cmp td{font-family:var(--mono);font-variant-numeric:tabular-nums}
.lt-cmp td.win{color:var(--good);font-weight:700}
.lt-cmp td.win::after{content:" \\25C2";letter-spacing:-.1em}
.lt-delta{font-size:.88rem;color:var(--ink);margin:.8rem 0 0;padding:.65rem .85rem;
  border-radius:9px;background:var(--surface-2);border:1px solid var(--line);line-height:1.6}
.lt-delta b{font-family:var(--mono)}
"""


# ── what to expect, before the clock starts ──────────────────────────────────────────
_EXPECT = {
    "measured": ("Measured end to end", "scored against the panel's own database"),
    "extracted": ("Provisions extracted", "runs live; accuracy not yet scored"),
    "reachable": ("Portal answers", "no adapter yet — discovery may return little or nothing"),
    "declared": ("Declared only", "the portal has not answered us; expect an empty run"),
}


def _expect_html(code: str, pillar: int, readiness: dict) -> str:
    """What this exact (economy, pillar) pair is likely to do.

    Written before the run rather than explained after it. An empty result from an economy at
    the "declared" rung and an empty result from a measured one look identical in the output
    and mean completely different things, and the person watching has sixty minutes.
    """
    row = (readiness or {}).get(code, {})
    level = row.get("level", "declared")
    head, why = _EXPECT.get(level, _EXPECT["declared"])
    measured = pillar in MEASURED_PILLARS
    pill = (f"pillar {pillar} definitions are scored against the answer key" if measured
            else f"pillar {pillar} definitions are coded from the Methodology, never scored")
    return (f'<div class="lt-slot"><h4>What to expect</h4>'
            f'<dl class="lt-kv">'
            f'<dt>Economy</dt><dd>{head} &mdash; {why}</dd>'
            f'<dt>Portal</dt><dd>{row.get("portal", "—")}</dd>'
            f'<dt>Pillar</dt><dd>{pill}</dd>'
            f'<dt>Blocker</dt><dd>{row.get("blocker", "—")}</dd>'
            f'</dl></div>')


# ── artefacts ────────────────────────────────────────────────────────────────────────
def _q(text: str) -> str:
    return str(text or "").replace('"', "'")


def run_record(state: dict) -> str:
    """The Run Record sheet, as CSV. Every field is read from the run, never typed."""
    b = state.get("brief", {})
    lines = ['"Team","Team VeriTrade"',
             f'"Task as read out","{_q(b.get("task"))}"',
             f'"Economy","{_q(b.get("economy"))}"',
             f'"Pillar","{_q(b.get("pillar"))}"',
             "",
             "Engine,Provider / model,Start (UTC),End (UTC),Elapsed (min),Cost (US$),"
             "Documents discovered,Documents fetched,Provisions,Rows exported"]
    for slot in ("A", "B"):
        r = state["runs"].get(slot)
        if not r:
            lines.append(f"Engine {slot},,,,,,,,,")
            continue
        lines.append(",".join([
            f"Engine {slot}", r["model"], r["start"], r["end"],
            f"{r['elapsed_min']:.1f}", f"{r['cost_usd']:.4f}",
            str(r["documents"]), str(r["fetched"]), str(r["provisions"]), str(r["rows"])]))
    return "\n".join(lines) + "\n"


def _diff(a: dict | None, b: dict | None, key: str, fmt: str) -> str:
    if not a or not b:
        return ""
    try:
        return fmt.format(b[key] - a[key])
    except (KeyError, TypeError):
        return ""


def engine_comparison(state: dict) -> str:
    """Engine Comparison, as CSV — the template's own rows, plus the difference it asks for."""
    a, b = state["runs"].get("A"), state["runs"].get("B")

    def cell(r, key, fmt="{}"):
        return fmt.format(r[key]) if r else ""

    rows = [("Field", "Engine A — first pass", "Engine B — second pass", "Difference (B − A)"),
            ("Provider / model", cell(a, "model"), cell(b, "model"), ""),
            ("Provisions exported", cell(a, "provisions"), cell(b, "provisions"),
             _diff(a, b, "provisions", "{:+d}")),
            ("Absent from the 2025 baseline", cell(a, "new"), cell(b, "new"),
             _diff(a, b, "new", "{:+d}")),
            ("Documents fetched during this pass", cell(a, "fetched"), cell(b, "fetched"),
             "second pass must be 0"),
            ("Elapsed (minutes)", cell(a, "elapsed_min", "{:.1f}"),
             cell(b, "elapsed_min", "{:.1f}"), _diff(a, b, "elapsed_min", "{:+.1f}")),
            ("Cost of this pass (US$)", cell(a, "cost_usd", "{:.4f}"),
             cell(b, "cost_usd", "{:.4f}"), _diff(a, b, "cost_usd", "{:+.4f}")),
            ("Rows exported", cell(a, "rows"), cell(b, "rows"), _diff(a, b, "rows", "{:+d}")),
            ("Rows needing review", cell(a, "review"), cell(b, "review"),
             _diff(a, b, "review", "{:+d}")),
            ("Rows the other engine did not find",
             str(a.get("only", "")) if a else "", str(b.get("only", "")) if b else "", "")]
    return "\n".join(",".join(f'"{_q(c)}"' for c in r) for r in rows) + "\n"


def short_note(state: dict) -> str:
    """The short note, in the organisers' own section order.

    Their template is a Word form with seven numbered parts and a signature block. Writing our
    own headings would make a steward hunt for each answer, so these are theirs — including
    section 6's checkbox, which we can honestly tick as *nothing was typed in by hand* because
    every figure above came off `RunMeta`.
    """
    b = state.get("brief", {})
    n = state.get("notes", {})
    a, bb = state["runs"].get("A"), state["runs"].get("B")
    eng = state.get("engines", {})

    def line(r, label):
        if not r:
            return f"- {label}: —"
        return (f"- {label}: {r['model']} — {r['provisions']} provisions, {r['rows']} rows, "
                f"{r['elapsed_min']:.1f} min, ${r['cost_usd']:.4f}, "
                f"{r['fetched']} documents fetched")

    by_hand = (n.get("by_hand") or "").strip()
    isolated = bool(bb) and bb.get("fetched") == 0
    out = [
        "# Live test — short note",
        "Finale morning, 15 October 2026 · Team VeriTrade", "",
        "## 1 · The run", "",
        f"**The task as read out:** {b.get('task') or '—'}",
        f"**Economy:** {b.get('economy', '—')}  ·  **Pillar:** {b.get('pillar', '—')}",
        f"**Engine A, first pass:** {eng.get('A', {}).get('provider', '—')} · "
        f"{eng.get('A', {}).get('model', '—')}",
        f"**Engine B, second pass:** {eng.get('B', {}).get('provider', '—')} · "
        f"{eng.get('B', {}).get('model', '—')}",
        f"**Submitted:** {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}", "",
        "## 2 · What came out", "",
        line(a, "Engine A"),
        line(bb, "Engine B"), "",
        ("Engine B fetched nothing: it re-read the documents engine A had already retrieved, "
         "so the only variable between the two passes is the engine." if isolated else
         "Engine B has not run, or fetched documents of its own — in which case the "
         "comparison is not engine-isolated and should be read with that in mind."), "",
        "## 3 · What worked", "",
        n.get("worked") or "_—_", "",
        "## 4 · What broke", "",
        n.get("broke") or "_—_", "",
        "## 5 · What a reviewer should be cautious about", "",
        n.get("caution") or "_—_", "",
        "## 6 · Anything done by hand", "",
        ("- [x] Nothing was typed in by hand. Every figure above is read from the run itself "
         "(`backend/metering.py` and `RunMeta`)." if not by_hand else
         "- [ ] Nothing was typed in by hand.\n"
         f"- [x] Something was: {by_hand}"), "",
        "## 7 · Declaration", "",
        "Everything submitted is my team's own work, produced by the system frozen at our "
        "declared release tag, using only the engines declared on 30 September.", "",
    ]
    return "\n".join(out) + "\n"


def short_note_docx(md: str) -> bytes | None:
    """The note as .docx, or None when python-docx is absent — the caller then offers the
    Markdown, so a missing optional dependency costs formatting and never the deliverable."""
    try:
        from docx import Document
    except Exception:                                        # noqa: BLE001
        return None
    doc = Document()
    for raw in md.splitlines():
        s = raw.strip()
        if not s:
            continue
        if s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("- "):
            doc.add_paragraph(s[2:].replace("**", ""), style="List Bullet")
        else:
            doc.add_paragraph(s.replace("**", ""))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── capture ──────────────────────────────────────────────────────────────────────────
def capture(state: dict, slot: str, result, started: str, finished: str) -> None:
    """Record one engine's pass. Reads the run; asks the operator for nothing.

    The whole RunResult is kept, not just its numbers: the evidence files are generated from it
    on the hand-in step, so what the operator downloads is the run in front of them rather than
    whatever a shared output directory happens to hold.
    """
    meta = result.meta
    rows = [m for m in result.mappings if m.law_name != "No provision found"]
    state["runs"][slot] = {
        "model": meta.model_version or meta.llm_provider,
        "provider": meta.llm_provider,
        "run_id": meta.run_id,
        "start": started, "end": finished,
        "elapsed_s": meta.processing_time_seconds,
        "elapsed_min": meta.processing_time_seconds / 60.0,
        "cost_usd": float((meta.cost or {}).get("total_usd", 0.0)),
        "documents": meta.docs_discovered,
        "fetched": meta.docs_fetched,
        "provisions": meta.provisions_extracted,
        "rows": len(rows),
        "review": sum(1 for m in rows if m.review_status.value == "pending_review"),
        "new": sum(1 for m in rows if m.discovery_tag.value == "NEW"),
        "keys": {f"{m.indicator_id}|{m.law_name}|{m.article_section}" for m in rows},
        "result": result,
    }
    a, b = state["runs"].get("A"), state["runs"].get("B")
    if a and b:                      # each engine's unique finds, computed not counted by hand
        a["only"] = len(a["keys"] - b["keys"])
        b["only"] = len(b["keys"] - a["keys"])


def _slot_html(slot: str, engine: dict, r: dict | None) -> str:
    head = (f'<h4>Engine {slot}<span class="eng">{engine.get("provider", "—")} · '
            f'{engine.get("model", "—")}</span></h4>')
    if not r:
        note = ("live — discovers and fetches" if slot == "A"
                else "second pass — re-reads engine A's documents, fetches nothing")
        return (f'<div class="lt-slot">{head}'
                f'<div class="lt-empty">Not run yet · {note}</div></div>')
    return (f'<div class="lt-slot">{head}<dl class="lt-kv">'
            f'<dt>Provisions</dt><dd>{r["provisions"]}</dd>'
            f'<dt>Rows exported</dt><dd>{r["rows"]}</dd>'
            f'<dt>Documents fetched</dt><dd>{r["fetched"]}</dd>'
            f'<dt>Elapsed</dt><dd>{r["elapsed_min"]:.1f} min</dd>'
            f'<dt>Cost</dt><dd>${r["cost_usd"]:.4f}</dd>'
            f'</dl></div>')


def _comparison_html(a: dict, b: dict) -> str:
    rows = [
        ("Provisions exported", a["provisions"], b["provisions"], "{}"),
        ("Absent from the 2025 baseline", a["new"], b["new"], "{}"),
        ("Documents fetched during this pass", a["fetched"], b["fetched"], "{}"),
        ("Elapsed (minutes)", a["elapsed_min"], b["elapsed_min"], "{:.1f}"),
        ("Cost of this pass (US$)", a["cost_usd"], b["cost_usd"], "{:.4f}"),
        ("Rows exported", a["rows"], b["rows"], "{}"),
        ("Rows needing review", a["review"], b["review"], "{}"),
        ("Found only by this engine", a.get("only", 0), b.get("only", 0), "{}"),
    ]
    body = [f'<tr><td>Provider / model</td><td>{a["model"]}</td><td>{b["model"]}</td></tr>']
    for label, va, vb, fmt in rows:
        want = BETTER.get(label)
        cls_a = cls_b = ""
        if want and va != vb:
            better_is_a = (va > vb) if want == "high" else (va < vb)
            cls_a, cls_b = (" class='win'", "") if better_is_a else ("", " class='win'")
        body.append(f'<tr><td>{label}</td><td{cls_a}>{fmt.format(va)}</td>'
                    f'<td{cls_b}>{fmt.format(vb)}</td></tr>')
    return ('<table class="lt-cmp"><thead><tr><th></th>'
            '<th>Engine A · first pass</th><th>Engine B · second pass</th></tr></thead>'
            f'<tbody>{"".join(body)}</tbody></table>')


def _delta_html(a: dict, b: dict) -> str:
    """The sentence the organisers asked for by name: the cost difference between the engines."""
    d = b["cost_usd"] - a["cost_usd"]
    rel = (f" ({abs(d) / a['cost_usd'] * 100:.0f}% {'more' if d > 0 else 'less'})"
           if a["cost_usd"] > 0 and d else "")
    speed = b["elapsed_min"] - a["elapsed_min"]
    cheaper = "engine B" if d < 0 else ("engine A" if d > 0 else "neither — identical")
    warn = ("" if b["fetched"] == 0 else
            f' <b style="color:var(--bad)">Engine B fetched {b["fetched"]} documents, so this '
            f'pass was not engine-isolated — the template requires 0.</b>')
    return (f'<p class="lt-delta">Engine B cost <b>${abs(d):.4f}</b> '
            f'{"less" if d < 0 else "more"} than engine A{rel} and took '
            f'<b>{abs(speed):.1f} min</b> {"less" if speed < 0 else "more"}. '
            f'Cheaper: <b>{cheaper}</b>.{warn}</p>')


# ── the surface ──────────────────────────────────────────────────────────────────────
def _engine_picker(state: dict, slot: str) -> None:
    """The UI-driven engine swap the criteria require.

    Defaults come from the declaration in `backend/config.py`, the same pair the README prints.
    Editable, because "engine swap should be UI-driven, not code-level" — and because on the
    day a declared model can be rate-limited and the steward needs to watch the switch happen
    rather than watch someone edit a file.
    """
    eng = state["engines"][slot]
    cols = st.columns([1, 1.7])
    with cols[0]:
        providers = list(reg.LLM_PROVIDERS)
        idx = providers.index(eng["provider"]) if eng["provider"] in providers else 0
        eng["provider"] = st.selectbox(f"Engine {slot} — provider", providers, index=idx,
                                       key=f"lt_prov_{slot}")
    with cols[1]:
        eng["model"] = st.text_input(f"Engine {slot} — model", value=eng["model"],
                                     key=f"lt_model_{slot}")


def render(state: dict, economies: dict[str, str],
           readiness: dict | None = None) -> dict | None:
    """Draw the surface. Returns a run request when the operator presses a run button."""
    st.markdown(f"<style>{CSS}</style>", unsafe_allow_html=True)
    state.setdefault("engines", new_state()["engines"])
    state.setdefault("notes", new_state()["notes"])
    step = state["step"]
    st.markdown(_steps_html(step), unsafe_allow_html=True)
    request = None
    readiness = readiness or {}

    if step == 0:
        st.caption("Type in whatever the steward reads out. Every economy and every pillar is "
                   "selectable — nothing here is prepared in advance.")
        order = economy_order(list(economies))
        c1, c2 = st.columns([2, 1])
        econ = c1.selectbox(
            "Economy", order, key="lt_econ",
            format_func=lambda c: economies.get(c, c)
            + ("" if c in LIVE_TEST_ORDER else "  (outside the nine)"))
        pillar = c2.selectbox("Pillar", list(range(1, 13)), index=5, key="lt_pillar",
                              format_func=lambda n: f"{n} · {PILLAR_SHORT.get(n, '')}")
        task = st.text_input("The task as read out", key="lt_task",
                             placeholder="the steward's words — this goes into the short note")
        st.markdown(_expect_html(econ, pillar, readiness), unsafe_allow_html=True)

        st.markdown("**The two declared engines** — fixed at submission on 30 September. Swap "
                    "them here if one is unreachable on the day; the run record and the note "
                    "report whatever was actually used.")
        _engine_picker(state, "A")
        _engine_picker(state, "B")

        if st.button("Start the clock", type="primary", width="stretch"):
            state["brief"] = {"economy": economies.get(econ, econ), "code": econ,
                              "pillar": pillar, "task": task}
            state["started"] = datetime.now(timezone.utc).isoformat(timespec="seconds")
            state["step"] = 1
            st.rerun()

    elif step == 1:
        b = state["brief"]
        st.caption(f"{b['economy']} · pillar {b['pillar']}"
                   + (f" · “{b['task']}”" if b.get("task") else ""))
        c1, c2 = st.columns(2)
        for col, slot in ((c1, "A"), (c2, "B")):
            with col:
                st.markdown(_slot_html(slot, state["engines"][slot], state["runs"].get(slot)),
                            unsafe_allow_html=True)
                done = slot in state["runs"]
                blocked = slot == "B" and "A" not in state["runs"]
                if st.button("Run again" if done else f"Run engine {slot}",
                             key=f"lt_run_{slot}", width="stretch", disabled=blocked,
                             help=("Engine A has to run first — engine B re-reads its "
                                   "documents" if blocked else None),
                             type="primary" if (slot == "A" and not done) else "secondary"):
                    request = {"slot": slot, "code": b["code"], "pillar": b["pillar"],
                               **state["engines"][slot]}
        if state["runs"] and st.button("See the comparison", type="primary", width="stretch"):
            state["step"] = 2
            st.rerun()

    elif step == 2:
        a, bb = state["runs"].get("A"), state["runs"].get("B")
        if a and bb:
            st.markdown(_comparison_html(a, bb), unsafe_allow_html=True)
            st.markdown(_delta_html(a, bb), unsafe_allow_html=True)
        elif a:
            st.info("Only engine A has run. The comparison — and criterion C5b — needs both.")
            m = st.columns(4)
            m[0].metric("Rows", a["rows"])
            m[1].metric("Provisions", a["provisions"])
            m[2].metric("Minutes", f"{a['elapsed_min']:.1f}")
            m[3].metric("Cost (US$)", f"{a['cost_usd']:.4f}")
        else:
            st.info("No run has been captured yet — go back a step and press Run.")
        if st.button("Prepare the hand-in", type="primary", width="stretch"):
            state["step"] = 3
            st.rerun()

    elif step == 3:
        _hand_in(state)

    if step and st.button("← Back a step", key="lt_back"):
        state["step"] = step - 1
        st.rerun()
    return request


def _hand_in(state: dict) -> None:
    """Everything that leaves the building, generated from the runs in front of the operator."""
    import json                                               # noqa: PLC0415

    from backend.export import csv_text                       # noqa: PLC0415
    from backend.export.json_export import build_payload      # noqa: PLC0415

    b = state.get("brief", {})
    stem = f"VeriTrade_{b.get('code', 'XX')}_P{b.get('pillar', '')}"

    st.markdown("**The evidence** — the mapped provisions themselves, in the official "
                "14-column format, one file per engine, plus the full JSON trace.")
    for slot in ("A", "B"):
        r = state["runs"].get(slot)
        if not r:
            continue
        res = r.get("result")
        cols = st.columns([1.6, 1, 1])
        cols[0].markdown(
            f'<div class="lt-slot" style="margin:0"><h4>Engine {slot}'
            f'<span class="eng">{r["model"]} · {r["rows"]} rows</span></h4></div>',
            unsafe_allow_html=True)
        if res is None:
            cols[1].caption("this run is not held in the session any more")
            continue
        cols[1].download_button(f"Evidence {slot} (.csv)", csv_text(res.mappings),
                                f"{stem}_engine{slot}.csv", "text/csv", width="stretch",
                                key=f"lt_csv_{slot}")
        cols[2].download_button(f"Trace {slot} (.json)",
                                json.dumps(build_payload(res), indent=2, ensure_ascii=False),
                                f"{stem}_engine{slot}.json", "application/json",
                                width="stretch", key=f"lt_json_{slot}")

    st.markdown("**The paperwork** — sections 3 to 6 are the only things on this screen a "
                "human writes. Every number came off the run.")
    n = state["notes"]
    c1, c2 = st.columns(2)
    n["worked"] = c1.text_area("3 · What worked", value=n.get("worked", ""), height=90,
                               placeholder="the part of the system you would show again")
    n["broke"] = c2.text_area("4 · What broke", value=n.get("broke", ""), height=90,
                              placeholder="the part you would not")
    n["caution"] = st.text_area("5 · What a reviewer should be cautious about",
                                value=n.get("caution", ""), height=70,
                                placeholder="name the rows or the indicator, not a general "
                                            "caution")
    n["by_hand"] = st.text_input("6 · Anything done by hand (leave blank if nothing was)",
                                 value=n.get("by_hand", ""))

    note_md = short_note(state)
    d1, d2, d3 = st.columns(3)
    d1.download_button("Run record (.csv)", run_record(state), f"{stem}_run_record.csv",
                       "text/csv", width="stretch")
    d2.download_button("Engine comparison (.csv)", engine_comparison(state),
                       f"{stem}_engine_comparison.csv", "text/csv", width="stretch")
    docx = short_note_docx(note_md)
    if docx:
        d3.download_button(
            "Short note (.docx)", docx, f"{stem}_short_note.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            width="stretch")
    else:
        d3.download_button("Short note (.md)", note_md, f"{stem}_short_note.md",
                           "text/markdown", width="stretch")
    with st.expander("Read the note before sending it"):
        st.markdown(note_md)


def brief_screen(*, economy: str, pillar: int, ocr_label: str, llm_label: str) -> None:
    """The live-test surface as a top-level screen, reached before any run exists.

    It used to be a tab inside the results, which meant it was unreachable until a run had
    already been done — and on the day there is no earlier run to open it from. Worse, the tab
    version ended with "run engine A from the sidebar", and the sidebar had been removed two
    redesigns earlier: an instruction pointing at a control that no longer exists.

    Here the run buttons run. Pressing one hands the request to the same pipeline the Run
    screen uses, records which engine slot it belongs to, and returns to this screen with the
    result captured — so the operator never leaves the checklist.
    """
    from . import geo                                   # noqa: PLC0415 — avoids an import cycle

    st.markdown("#### Live stress test")
    st.caption("One economy and one pillar, read out by the steward on the day. Engine A runs "
               "live; engine B re-reads engine A's documents and fetches nothing, so the only "
               "difference between the two passes is the engine.")
    if "livetest" not in st.session_state:
        st.session_state["livetest"] = new_state()
    econ_names = st.session_state.get("_econ_names") or {}
    request = render(st.session_state["livetest"], econ_names, readiness=geo.readiness())
    if request:
        # Drive the ordinary pipeline. The slot is remembered so the completed run is filed
        # against the right engine without the operator copying anything.
        st.session_state["economy"] = request["code"]
        st.session_state["pillar"] = request["pillar"]
        st.session_state["use_samples"] = False
        st.session_state["fresh_run"] = True
        st.session_state["llm_provider"] = request["provider"]
        st.session_state["llm_model"] = request["model"]
        st.session_state["lt_pending"] = request["slot"]
        st.session_state["lt_started"] = datetime.now(timezone.utc).isoformat(timespec="seconds")
        # The whole point of the second pass: no portal is contacted.
        first = st.session_state["livetest"]["runs"].get("A", {}).get("result")
        st.session_state["lt_reuse"] = (list(first.meta.documents)
                                        if request["slot"] == "B" and first else None)
        st.session_state["run_requested"] = True
        st.rerun()
